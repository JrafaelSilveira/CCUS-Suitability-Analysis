"""Generate the 'Interactive map and analysis program' section of the report.

American English, formal register, prose-led. Targeted at no more than two
pages so that the section does not exceed the length of the surrounding
project report. Two figure placeholders for screenshots.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_figure_placeholder(doc, caption_no, caption_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.rows[0].cells[0]
    cell.text = f"[ Insert Figure {caption_no} — screenshot ]"
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figure {caption_no}: {caption_text}")
    cap_run.font.size = Pt(10)
    cap_run.font.italic = True


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "3.1  Interactive map and analysis program", level=1)

    add_paragraph(doc,
        "A dedicated desktop program, CCUS Suitability Analysis, was "
        "developed alongside the present study to make the screening of "
        "candidate sites for in-situ CO₂ mineralization both interactive "
        "and reproducible. The program is implemented in Python with a "
        "PySide6 (Qt 6) graphical interface and runs as a native Windows "
        "application; no browser or web server is required. Its purpose "
        "is to expose the analysis defined in this report through a "
        "graphical environment that geoscientists can operate directly, "
        "while keeping the underlying scientific code available for review "
        "and re-use."
    )

    add_figure_placeholder(doc, caption_no="3.1",
        caption_text=(
            "Home screen of CCUS Suitability Analysis. The user enters the "
            "analysis environment, loads an alternative geodatabase, or "
            "consults the About page from this view."
        ),
    )

    add_paragraph(doc,
        "The interface is organized into four tabs over a single dataset. "
        "A dashboard presents the multi-criteria weighting and a tabulated "
        "ranking of the candidate polygons; a three-dimensional viewer "
        "drapes the bedrock polygons over the digital elevation model; an "
        "interactive two-dimensional map renders the capacity heatmap "
        "directly inside the application; and a final tab collects the "
        "static figures referenced elsewhere in this report. The "
        "frequently used controls — running the analysis, loading a "
        "different geodatabase, opening the About page and the Settings "
        "dialog — are gathered in a toolbar at the top of the window."
    )

    add_paragraph(doc,
        "The backend reads vector data from an Esri File Geodatabase by "
        "way of GeoPandas and the pyogrio driver, defaulting to the NGU "
        "BerggrunnN250 bedrock map at 1:250 000. Each polygon is "
        "reprojected into a metric coordinate system, classified into a "
        "rock family, and assigned a storage capacity from the "
        "volumetric formulation V = A · h · φ · E · ρ_CO₂, in which A is "
        "the polygon area, h the assumed thickness of the fracture zone "
        "(200 m by default), φ the effective porosity (1.5 % by default, "
        "overridden where measured values are available), E the storage "
        "efficiency factor (2 %), and ρ_CO₂ ≈ 700 kg/m³ at supercritical "
        "conditions. The estimated mass is then converted to megatonnes "
        "and reported both in the dashboard and as a colour scale on the "
        "interactive map."
    )

    add_paragraph(doc,
        "Two suitability metrics operate on top of the capacity calculation. "
        "A weighted linear combination of reservoir quality, fault and "
        "injectivity behaviour, structural setting and petrophysical "
        "support produces a single score whose weights are exposed as "
        "sliders in the dashboard. An alternative analytic hierarchy "
        "process derives a consistent weight vector from pairwise "
        "comparisons of the same four criteria. A seal or cap rock "
        "criterion is intentionally omitted, since in-situ mineralization "
        "fixes the injected CO₂ in a chemically stable carbonate form. "
        "Groundwater boreholes from the NGU Grunnvannsborehull dataset "
        "are loaded as a context layer on the capacity map, but do not "
        "enter the formula above: their reported yield measures "
        "permeability rather than porosity, and most wells are shallower "
        "than the supercritical-CO₂ depth window targeted here."
    )

    add_figure_placeholder(doc, caption_no="3.2",
        caption_text=(
            "Interactive capacity heatmap embedded inside the program. "
            "Warm colours mark the polygons with the highest estimated "
            "storage capacity; NGU groundwater boreholes are overlaid as "
            "togglable context layers."
        ),
    )

    add_paragraph(doc,
        "In typical use the analyst loads the relevant geodatabase, "
        "presses Run Analysis on the toolbar, and inspects the resulting "
        "ranking and maps. The weighting can then be revised on the "
        "dashboard and the ranking recomputed in place; alternative "
        "renderings, including the WLC and AHP score maps and the basic "
        "rock-family view, are accessible from the same tab. The "
        "interactive map can be expanded to fill the application pane "
        "when a closer inspection is needed. With the exception of the "
        "basemap tiles retrieved from CartoDB, the program runs entirely "
        "on the user's machine, preserving the interactivity of the "
        "underlying Jupyter notebook while keeping the input data local."
    )

    out = "Program_section.docx"
    try:
        doc.save(out)
    except PermissionError:
        # User has the old version open in Word; write to a versioned name.
        out = "Program_section_v2.docx"
        doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build()
